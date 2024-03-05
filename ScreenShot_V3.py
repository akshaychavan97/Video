import sys
import os
import threading

import keyboard
import pyautogui
import logging

from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QLineEdit, QLabel, QTextEdit,QCheckBox
from docx import Document
from docx.shared import Inches
from PyQt5.QtCore import Qt


# @Author - Akshay Chavan - akshay.chavan@voya.com - i723130
# @Author - goutham S - goutham.satheesha@voya.com - i732746

# Configure logging
logging.basicConfig(filename='log_ScreenshotApp.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')


class ScreenshotApp(QMainWindow):
    def __init__(self):
        super().__init__()

        self.take_screenshot_flag = False
        self.screenshot_counter = 1
        self.version_number = 1
        self.test_case_name = "Evidence"  # Default value for Test Case Name
        self.screenshot_key = "home"  # Default value for Enter Key for Screenshot (Qt.Key_Q corresponds to "q")
        self.document_name=""
        self.initUI()
        # Start the keypress detection thread
        self.keypress_thread = threading.Thread(target=self.detect_keypress)
        self.keypress_thread.daemon = True  # Allow the thread to be Terminated when the main program exits
        self.keypress_thread.start()

    def initUI(self):
        self.setGeometry(100, 100, 600, 200)
        self.setWindowTitle("Screenshot App")
        font = QFont()
        font.setBold(True)


        # Set background color
        self.setStyleSheet("background-color: lightblue;")

        self.test_case_label = QLineEdit(self)
        self.test_case_label.setPlaceholderText("Enter Test Case Name (Default: Evidence)")
        self.test_case_label.setGeometry(50, 20, 500, 30)
        self.test_case_label.setStyleSheet("background-colour: white;")

        self.output_location = QLineEdit(self)
        self.output_location.setPlaceholderText("Enter Output File Path (Default: Output Folder)")
        self.output_location.setGeometry(50, 55, 500, 30)
        self.output_location.setStyleSheet("background-colour: white;")

        self.key_label = QLabel("Enter Key for Screenshot (Default: home):", self)
        self.key_label.setGeometry(50, 90, 250, 30)
        self.test_case_label.setStyleSheet("background-colour: white;")

        self.key_label = QLabel("Press END Key or Button to Create Word Document", self)
        self.key_label.setGeometry(50, 125, 280, 30)
        self.test_case_label.setStyleSheet("background-colour: white;")

        self.key_input = QLineEdit(self)
        self.key_input.setText("home")  # Set the default key
        self.key_input.setGeometry(260, 90, 100, 30)
        self.key_input.setAlignment(Qt.AlignCenter)
        self.key_input.setFont(font)

        self.start_button = QPushButton("Start", self)
        self.start_button.setGeometry(50, 160, 150, 30)
        self.start_button.setStyleSheet("background-color: green; color: white;")
        self.start_button.clicked.connect(self.start_screenshot)

        self.end_button = QPushButton("End", self)
        self.end_button.setGeometry(220, 160, 150, 30)
        self.end_button.setStyleSheet("background-color: red; color: white;")
        self.end_button.clicked.connect(self.end_screenshot)
        self.end_button.setEnabled(False)

        self.update_checkbox = QCheckBox("Use Existing File",self)
        self.update_checkbox.setGeometry(400, 90, 150, 30)

        # Screenshot counter label
        self.counter_label = QLabel("Screenshot Counter: 0", self)
        self.counter_label.setGeometry(400, 125, 150, 30)

        self.developer_label = QLabel("@Developed By Voya India", self)
        self.developer_label.setGeometry(400, 170, 160, 30)




    def start_screenshot(self):
        self.document=Document()
        self.test_case_name = self.test_case_label.text().strip() or "Evidence"  # Use default if not entered
        self.screenshot_key = self.key_input.text().strip() or "home"  # Use default if not entered

        self.output_location_path = self.output_location.text().strip() or os.path.join(os.getcwd(), "Output")
        # Use default if not entered
        folder_name=self.output_location_path
        # Create a folder with just the test case name

        os.makedirs(folder_name, exist_ok=True)

        # Check if the document already exists
        if self.update_checkbox.isChecked():
            self.document_name = os.path.join(folder_name, f'{self.test_case_name}.docx')
        else:
            self.document_name = os.path.join(folder_name, f'{self.test_case_name}_v{self.version_number}.docx')
            while os.path.exists(os.path.join(folder_name, f'{self.test_case_name}_v{self.version_number}.docx')):
                self.version_number += 1
            # Append the version number to the document name
            self.document_name = os.path.join(folder_name, f'{self.test_case_name}_v{self.version_number}.docx')
        if os.path.exists(self.document_name):
            # Find a version number to add to the document name
            existing_dcoument=Document(self.document_name)
            self.screenshot_counter=len(existing_dcoument.inline_shapes)+1

        self.document.add_heading('Test Case: ' + self.test_case_name, 0)
        logging.info(f"Document Version: {self.version_number}")
        logging.info(f"Document Path: {self.document_name}")
        self.take_screenshot_flag = True
        self.start_button.setEnabled(False)
        self.end_button.setEnabled(True)
        self.start_button.setStyleSheet("background-color: lightblue; color:white")

        # Gray out the input fields
        self.test_case_label.setEnabled(False)
        self.key_input.setEnabled(False)

        # Log the start action
        logging.info(f"Test Case: {self.test_case_name} - Start Screenshot")
        self.showMinimized()

    def end_screenshot(self):
        self.take_screenshot_flag = False
        self.start_button.setEnabled(True)
        self.end_button.setEnabled(False)

        # Enable the input fields
        self.test_case_label.setEnabled(True)
        self.key_input.setEnabled(True)
        self.start_button.setStyleSheet("background-color: green; color: white;")
        if self.screenshot_counter > 1:
            self.document.save(self.document_name)

            # Log the document save action
            logging.info(f"Test Case: {self.test_case_name} - Document Saved")
            self.screenshot_counter = 1

        # Log the end action
        logging.info(f"Test Case: {self.test_case_name} - End Screenshot")

    def takescreenshot(self):
        logging.info(f"Taking ScreenShot")

        screenshot_name = os.path.join(os.getcwd(), "Output",
                                         f'{self.test_case_name}_screenshot_v{self.version_number}_{self.screenshot_counter}.png')
        pyautogui.screenshot(screenshot_name)
        default_description = "Description :- "  # Default description
        if self.screenshot_counter > 1:
            self.document.add_page_break()
        self.document.add_heading('Step ' + str(self.screenshot_counter), level=1)
        self.document.add_paragraph(default_description)
        self.document.add_paragraph('Screenshot No :- ' + str(self.screenshot_counter) )
        self.document.add_picture(screenshot_name, width=Inches(6))
        os.remove(screenshot_name)
        self.increment_counter()

    def increment_counter(self):
        self.counter_label.setText(f"Screenshot Counter: {self.screenshot_counter}")
        self.screenshot_counter += 1


    def detect_keypress(self):
        logging.info(f"Press {self.key_input} to Capture ScreenShot")
        while True:
            event = keyboard.read_event()
            if event.event_type == keyboard.KEY_DOWN and event.name == self.screenshot_key and self.take_screenshot_flag:
                self.takescreenshot()
            elif event.event_type == keyboard.KEY_DOWN and event.name == "end" and self.take_screenshot_flag:
                self.end_screenshot()


def main():
    app = QApplication(sys.argv)
    window = ScreenshotApp()
    window.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
